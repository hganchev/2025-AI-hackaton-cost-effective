"""
Book content extraction utilities for various file formats.
"""
import os
import re
import logging
import requests
from io import BytesIO
from urllib.parse import urlparse
from pathlib import Path

# File format specific imports
import PyPDF2
from ebooklib import epub
from bs4 import BeautifulSoup
import docx
import markdown

logger = logging.getLogger(__name__)

class BookExtractor:
    """
    Class responsible for extracting text content from books in various formats.
    """
    
    @staticmethod
    def extract_from_book(book):
        """
        Extract text from a book model based on its source (file or URL) and format.
        
        Args:
            book: Book model instance
            
        Returns:
            str: Extracted text content
        """
        try:
            if book.file:
                file_path = book.file.path
                file_format = book.file_format or Path(file_path).suffix.lstrip('.')
                
                # Use the appropriate extractor based on file format
                if file_format == 'pdf':
                    return BookExtractor.extract_from_pdf(file_path)
                elif file_format == 'epub':
                    return BookExtractor.extract_from_epub(file_path)
                elif file_format == 'txt':
                    return BookExtractor.extract_from_txt(file_path)
                elif file_format == 'docx':
                    return BookExtractor.extract_from_docx(file_path)
                elif file_format == 'html':
                    return BookExtractor.extract_from_html(file_path)
                elif file_format == 'md':
                    return BookExtractor.extract_from_markdown(file_path)
                else:
                    # Default to treating as text file
                    return BookExtractor.extract_from_txt(file_path)
            elif book.url:
                return BookExtractor.extract_from_url(book.url, book.file_format)
            else:
                return "No content source available for this book."
        except Exception as e:
            logger.error(f"Error extracting content from book {book.id}: {str(e)}")
            return f"Error extracting content: {str(e)}"
    
    @staticmethod
    def extract_from_pdf(file_path):
        """
        Extract text content from PDF files using PyPDF2.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            text = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                # Limit to first 50 pages for very large PDFs to prevent memory issues
                max_pages = min(num_pages, 50)
                
                for page_num in range(max_pages):
                    page = reader.pages[page_num]
                    text.append(page.extract_text())
                    
                # Add a note if we didn't process all pages
                if num_pages > max_pages:
                    text.append(f"\n[Note: Only the first {max_pages} pages of {num_pages} were processed]")
                    
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return f"Error extracting PDF content: {str(e)}"
    
    @staticmethod
    def extract_from_epub(file_path):
        """
        Extract text content from EPUB files using ebooklib.
        
        Args:
            file_path (str): Path to the EPUB file
            
        Returns:
            str: Extracted text content
        """
        try:
            book = epub.read_epub(file_path)
            chapters = []
            
            # Extract title and author info if available
            if book.get_metadata('DC', 'title'):
                title = book.get_metadata('DC', 'title')[0][0]
                chapters.append(f"Title: {title}")
                
            if book.get_metadata('DC', 'creator'):
                author = book.get_metadata('DC', 'creator')[0][0]
                chapters.append(f"Author: {author}")
                
            chapters.append("\n--- Content ---\n")
            
            # Extract content from all HTML documents in the EPUB
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    # Parse HTML content
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    # Extract text and clean up whitespace
                    text = soup.get_text(separator=' ', strip=True)
                    chapters.append(text)
                    
            return "\n\n".join(chapters)
        except Exception as e:
            logger.error(f"EPUB extraction error: {str(e)}")
            return f"Error extracting EPUB content: {str(e)}"
    
    @staticmethod
    def extract_from_txt(file_path):
        """
        Extract text content from plain text files with encoding handling.
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            str: Extracted text content
        """
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"TXT extraction error with {encoding}: {str(e)}")
                
        # If all encodings fail
        return f"Error: Unable to decode text file with supported encodings. The file may be in a different encoding or is not a valid text file."
    
    @staticmethod
    def extract_from_docx(file_path):
        """
        Extract text content from DOCX files using python-docx.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract document properties if available
            props = []
            try:
                core_props = doc.core_properties
                if core_props.title:
                    props.append(f"Title: {core_props.title}")
                if core_props.author:
                    props.append(f"Author: {core_props.author}")
                if props:
                    props.append("\n--- Content ---\n")
            except:
                pass
                
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    full_text.append(para.text)
                    
            # Include tables content
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():  # Skip empty rows
                        full_text.append(row_text)
                
            return "\n".join(props + full_text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return f"Error extracting DOCX content: {str(e)}"
    
    @staticmethod
    def extract_from_html(file_path):
        """
        Extract text content from HTML files using BeautifulSoup.
        
        Args:
            file_path (str): Path to the HTML file
            
        Returns:
            str: Extracted text content
        """
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                # Parse HTML content
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove script and style elements
                for script_or_style in soup(["script", "style"]):
                    script_or_style.extract()
                    
                # Get title if available
                title = soup.title.string if soup.title else None
                
                # Extract text
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up excessive whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                # Add title at the beginning if available
                if title:
                    text = f"Title: {title}\n\n{text}"
                    
                return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"HTML extraction error with {encoding}: {str(e)}")
                
        # If all encodings fail
        return f"Error: Unable to decode HTML file with supported encodings."
    
    @staticmethod
    def extract_from_markdown(file_path):
        """
        Extract text content from Markdown files.
        
        Args:
            file_path (str): Path to the Markdown file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Read the markdown file
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Option 1: Return markdown content directly
            # return md_content
            
            # Option 2: Convert markdown to HTML and extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text(separator='\n\n', strip=True)
            
        except UnicodeDecodeError:
            try:
                # Try with another common encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    md_content = f.read()
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text(separator='\n\n', strip=True)
            except Exception as e:
                logger.error(f"Markdown extraction error (latin-1): {str(e)}")
                return f"Error extracting Markdown content: {str(e)}"
        except Exception as e:
            logger.error(f"Markdown extraction error: {str(e)}")
            return f"Error extracting Markdown content: {str(e)}"
    
    @staticmethod
    def extract_from_url(url, format_hint=None):
        """
        Extract content from a URL by downloading and processing it.
        
        Args:
            url (str): URL to download content from
            format_hint (str, optional): Hint about the expected format
            
        Returns:
            str: Extracted text content
        """
        try:
            # Download content from URL
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Determine the content type
            content_type = response.headers.get('Content-Type', '').lower()
            
            # Create a temporary file-like object
            content = BytesIO(response.content)
            
            # Use format hint or determine from URL/content-type
            if format_hint:
                file_format = format_hint
            else:
                # Try to determine format from URL
                parsed_url = urlparse(url)
                path = parsed_url.path.lower()
                
                if path.endswith('.pdf'):
                    file_format = 'pdf'
                elif path.endswith('.epub'):
                    file_format = 'epub'
                elif path.endswith('.txt'):
                    file_format = 'txt'
                elif path.endswith('.docx'):
                    file_format = 'docx'
                elif path.endswith('.html') or path.endswith('.htm'):
                    file_format = 'html'
                elif path.endswith('.md'):
                    file_format = 'md'
                # Try to determine from content-type
                elif 'application/pdf' in content_type:
                    file_format = 'pdf'
                elif 'application/epub+zip' in content_type:
                    file_format = 'epub'
                elif 'text/plain' in content_type:
                    file_format = 'txt'
                elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                    file_format = 'docx'
                elif 'text/html' in content_type:
                    file_format = 'html'
                elif 'text/markdown' in content_type:
                    file_format = 'md'
                else:
                    # Default to HTML for web pages
                    file_format = 'html'
            
            # Extract content based on determined format
            if file_format == 'pdf':
                reader = PyPDF2.PdfReader(content)
                texts = [page.extract_text() for page in reader.pages]
                return "\n\n".join(texts)
            elif file_format == 'epub':
                # EPUBs can't be read directly from BytesIO
                # We would need to save to a temp file first
                return f"EPUB downloading not implemented: {url}"
            elif file_format == 'txt':
                return response.text
            elif file_format == 'docx':
                # DOCXs can't be read directly from BytesIO
                return f"DOCX downloading not implemented: {url}"
            elif file_format == 'html':
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Remove script and style elements
                for script_or_style in soup(["script", "style"]):
                    script_or_style.extract()
                    
                # Get title if available
                title = soup.title.string if soup.title else None
                
                # Extract text
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up excessive whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                # Add title at the beginning if available
                if title:
                    text = f"Title: {title}\n\n{text}"
                    
                return text
            elif file_format == 'md':
                html = markdown.markdown(response.text)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text(separator='\n\n', strip=True)
            else:
                return f"Unsupported content type for URL: {url} (Content-Type: {content_type})"
                
        except requests.RequestException as e:
            logger.error(f"URL download error: {str(e)}")
            return f"Error downloading content from URL: {str(e)}"
        except Exception as e:
            logger.error(f"URL content extraction error: {str(e)}")
            return f"Error extracting content from URL: {str(e)}"